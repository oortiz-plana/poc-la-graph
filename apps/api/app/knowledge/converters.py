"""Local, non-networked adapters for supported source formats."""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import PurePosixPath
from typing import Any, cast


class DocumentConversionError(ValueError):
    """A source cannot be converted safely into usable normalized text."""


@dataclass(frozen=True)
class ConvertedDocument:
    text: str
    media_type: str
    metadata: dict[str, Any]


MEDIA_TYPES = {
    ".md": "text/markdown",
    ".txt": "text/plain",
    ".html": "text/html",
    ".htm": "text/html",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def convert_document(
    relative_path: str,
    raw: bytes,
    *,
    max_extracted_bytes: int,
) -> ConvertedDocument:
    suffix = PurePosixPath(relative_path).suffix.lower()
    media_type = MEDIA_TYPES.get(suffix)
    if media_type is None:
        raise DocumentConversionError("Unsupported source extension")
    if suffix in {".md", ".txt"}:
        text = _strict_utf8(raw, relative_path)
        converter = "utf8-identity"
    elif suffix in {".html", ".htm"}:
        text = _html(raw, relative_path)
        converter = "haystack-html-local"
    elif suffix == ".pdf":
        text = _pdf(raw)
        converter = "haystack-pypdf-local"
    else:
        text = _docx(raw, max_extracted_bytes)
        converter = "haystack-docx-local"
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if not text.strip():
        raise DocumentConversionError("Document conversion produced no text")
    if len(text.encode("utf-8")) > max_extracted_bytes:
        raise DocumentConversionError(
            "Extracted document text exceeds configured limit"
        )
    return ConvertedDocument(
        text=text,
        media_type=media_type,
        metadata={"adapter": converter, "networkAccess": False, "version": 1},
    )


def _strict_utf8(raw: bytes, relative_path: str) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentConversionError(
            f"Knowledge document is not valid UTF-8: {relative_path}"
        ) from exc


def _html(raw: bytes, relative_path: str) -> str:
    text = _strict_utf8(raw, relative_path)
    if not re.search(r"<(?:!doctype\s+html|html|body|head|article|main)\b", text, re.I):
        raise DocumentConversionError("HTML source has no recognized HTML signature")
    try:
        from haystack.components.converters import HTMLToDocument
        from haystack.dataclasses import ByteStream

        result = HTMLToDocument(
            extraction_kwargs={
                "include_comments": False,
                "include_tables": True,
                "include_links": False,
                "output_format": "markdown",
                "no_fallback": False,
            },
            encoding="utf-8",
        ).run(
            sources=[
                ByteStream(
                    data=raw,
                    meta={"file_path": relative_path, "encoding": "utf-8"},
                    mime_type="text/html",
                )
            ]
        )
        documents = result["documents"]
        extracted = documents[0].content if documents else None
    except Exception as exc:
        raise DocumentConversionError("HTML conversion failed") from exc
    if extracted is None:
        raise DocumentConversionError("HTML conversion produced no text")
    return cast(str, extracted)


def _pdf(raw: bytes) -> str:
    if not raw.startswith(b"%PDF-"):
        raise DocumentConversionError("PDF source has an invalid signature")
    try:
        from haystack.components.converters import PyPDFToDocument
        from haystack.dataclasses import ByteStream
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw), strict=True)
        if reader.is_encrypted:
            raise DocumentConversionError("Encrypted PDFs are not supported")
        # Force strict parsing before the Haystack adapter performs normalized
        # page extraction and preserves form-feed boundaries.
        tuple(reader.pages)
        result = PyPDFToDocument().run(
            sources=[
                ByteStream(
                    data=raw,
                    meta={"file_path": "source.pdf"},
                    mime_type="application/pdf",
                )
            ]
        )
        documents = result["documents"]
        text = documents[0].content if documents else None
    except DocumentConversionError:
        raise
    except Exception as exc:
        raise DocumentConversionError("PDF conversion failed") from exc
    if not text or not text.strip():
        raise DocumentConversionError("PDF has no usable embedded text")
    return text


def _validate_docx_archive(raw: bytes, max_extracted_bytes: int) -> None:
    if not raw.startswith(b"PK\x03\x04"):
        raise DocumentConversionError("DOCX source has an invalid ZIP signature")
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            infos = archive.infolist()
            if len(infos) > 10_000:
                raise DocumentConversionError("DOCX archive has too many entries")
            total = 0
            names: set[str] = set()
            for info in infos:
                name = PurePosixPath(info.filename)
                if name.is_absolute() or ".." in name.parts:
                    raise DocumentConversionError("DOCX archive path is unsafe")
                if info.filename in names:
                    raise DocumentConversionError("DOCX archive has duplicate entries")
                names.add(info.filename)
                if info.flag_bits & 0x1:
                    raise DocumentConversionError("Encrypted DOCX entries are unsafe")
                total += info.file_size
                if total > max(max_extracted_bytes * 8, 16 * 1024 * 1024):
                    raise DocumentConversionError("DOCX archive expands beyond limit")
                if info.file_size > 1_000_000 and info.compress_size == 0:
                    raise DocumentConversionError("DOCX archive compression is invalid")
                if (
                    info.compress_size
                    and info.file_size / info.compress_size > 200
                    and info.file_size > 1_000_000
                ):
                    raise DocumentConversionError(
                        "DOCX archive compression ratio is unsafe"
                    )
                if info.filename.lower().endswith(".xml"):
                    with archive.open(info) as stream:
                        prefix = stream.read(65_536).upper()
                    if b"<!DOCTYPE" in prefix or b"<!ENTITY" in prefix:
                        raise DocumentConversionError(
                            "DOCX XML declarations are unsafe"
                        )
            if "word/document.xml" not in names or "[Content_Types].xml" not in names:
                raise DocumentConversionError("DOCX container is incomplete")
    except DocumentConversionError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise DocumentConversionError("DOCX container is invalid") from exc


def _docx(raw: bytes, max_extracted_bytes: int) -> str:
    _validate_docx_archive(raw, max_extracted_bytes)
    try:
        from docx import Document as DocxDocument
        from docx.document import Document as DocxDocumentType
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = DocxDocument(io.BytesIO(raw))
        blocks: list[str] = []
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = Paragraph(child, document)
                page_break = bool(
                    child.xpath(
                        './/*[local-name()="br" and @*[local-name()="type"]="page"]'
                    )
                    or child.xpath('.//*[local-name()="lastRenderedPageBreak"]')
                )
                value = paragraph.text.strip()
                if value:
                    style = paragraph.style.name if paragraph.style else ""
                    heading = re.match(r"Heading\s+([1-6])$", style, re.I)
                    blocks.append(
                        f"{'#' * int(heading.group(1))} {value}" if heading else value
                    )
                if page_break:
                    blocks.append("\f")
            elif child.tag.endswith("}tbl"):
                table = Table(child, document)
                rows = [
                    [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    for row in table.rows
                ]
                if rows:
                    width = max(len(row) for row in rows)
                    normalized = [row + [""] * (width - len(row)) for row in rows]
                    blocks.append("| " + " | ".join(normalized[0]) + " |")
                    blocks.append("| " + " | ".join(["---"] * width) + " |")
                    blocks.extend(
                        "| " + " | ".join(row) + " |" for row in normalized[1:]
                    )
        if not isinstance(document, DocxDocumentType):
            raise DocumentConversionError("DOCX conversion failed")
        return "\n\n".join(blocks).replace("\n\n\f\n\n", "\f")
    except DocumentConversionError:
        raise
    except Exception as exc:
        raise DocumentConversionError("DOCX conversion failed") from exc
