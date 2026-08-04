from __future__ import annotations

import hashlib
import io
import json
import sqlite3
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
import tiktoken
from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK
from pydantic import ValidationError
from pypdf import PdfWriter

from app.config.settings import Settings
from app.knowledge.chunking import build_chunks
from app.knowledge.profiles import parse_document_profiles
from app.knowledge.source_index import SourceIndex
from app.knowledge.sources import FilesystemDocumentSource, SourceValidationError


def profiles(
    *,
    leaf: int = 64,
    parent: int = 128,
    overlap: int = 8,
) -> dict[str, object]:
    return {
        "version": 1,
        "defaultProfile": "generic",
        "rules": [
            {"glob": "priority-*.md", "profile": "first"},
            {"glob": "*.md", "profile": "second"},
        ],
        "profiles": {
            name: {
                "hardBoundaries": [
                    "page",
                    "legal_article",
                    "markdown_heading",
                ],
                "leafTokens": leaf,
                "parentTokens": parent,
                "overlapTokens": overlap,
                "autoMergeThreshold": 0.5,
            }
            for name in ("generic", "first", "second")
        },
    }


def text_pdf(*pages: str) -> bytes:
    kids = " ".join(f"{3 + index} 0 R" for index in range(len(pages)))
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        (f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>").encode(),
    ]
    font_id = 3 + len(pages)
    content_start = font_id + 1
    for index in range(len(pages)):
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
                f"/Contents {content_start + index} 0 R >>"
            ).encode()
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for page in pages:
        escaped = page.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )
    value = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, item in enumerate(objects, start=1):
        offsets.append(len(value))
        value.extend(f"{number} 0 obj\n".encode() + item + b"\nendobj\n")
    xref = len(value)
    value.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    value.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        value.extend(f"{offset:010d} 00000 n \n".encode())
    value.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n"
        ).encode()
    )
    return bytes(value)


def test_profile_rules_are_ordered_and_fingerprint_tracks_configuration() -> None:
    original = parse_document_profiles(json.dumps(profiles()))
    changed = parse_document_profiles(json.dumps(profiles(overlap=9)))

    assert original.select("priority-law.md")[0] == "first"
    assert original.select("other.md")[0] == "second"
    assert original.select("other.txt")[0] == "generic"
    assert original.processing_fingerprint() != changed.processing_fingerprint()


@pytest.mark.parametrize(
    "mutation",
    [
        lambda value: value.update(defaultProfile="missing"),
        lambda value: value["rules"].append(
            {"glob": "priority-*.md", "profile": "second"}
        ),
        lambda value: value["profiles"]["generic"].update(leafTokens=32),
        lambda value: value["profiles"]["generic"].update(parentTokens=64),
        lambda value: value["profiles"]["generic"].update(overlapTokens=64),
        lambda value: value["profiles"]["generic"].update(hardBoundaries=["arbitrary"]),
    ],
)
def test_invalid_profiles_are_rejected_at_settings_startup(mutation: object) -> None:
    value = profiles()
    mutation(value)  # type: ignore[operator]
    with pytest.raises((ValueError, ValidationError)):
        Settings(knowledge_document_profiles_json=json.dumps(value))


def test_mixed_text_html_and_docx_discovery_preserves_raw_bytes(
    tmp_path: Path,
) -> None:
    (tmp_path / "a.md").write_bytes(b"# Heading\nMarkdown")
    (tmp_path / "b.txt").write_bytes(b"Plain text")
    (tmp_path / "c.html").write_bytes(
        b"<!doctype html><html><body><main><h1>Title</h1>"
        b"<p>Local HTML body.</p></main></body></html>"
    )
    (tmp_path / "ignored.rtf").write_bytes(b"unsupported")
    docx_path = tmp_path / "d.docx"
    document = DocxDocument()
    document.add_heading("DOCX heading", level=1)
    document.add_paragraph("Before break").add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("After break")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    document.save(docx_path)

    snapshot = FilesystemDocumentSource(tmp_path).discover()
    by_path = {item.relative_path: item for item in snapshot.documents}

    assert list(by_path) == ["a.md", "b.txt", "c.html", "d.docx"]
    assert by_path["b.txt"].content == "Plain text"
    assert "Local HTML body" in by_path["c.html"].content
    assert "# DOCX heading" in by_path["d.docx"].content
    assert "| A | B |" in by_path["d.docx"].content
    assert "\f" in by_path["d.docx"].content
    for path, item in by_path.items():
        raw = (tmp_path / path).read_bytes()
        assert item.raw_bytes == raw
        assert item.sha256 == hashlib.sha256(raw).hexdigest()


def test_pdf_signature_empty_text_and_invalid_docx_are_rejected(
    tmp_path: Path,
) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    stream = io.BytesIO()
    writer.write(stream)
    (tmp_path / "empty.pdf").write_bytes(stream.getvalue())
    with pytest.raises(SourceValidationError):
        FilesystemDocumentSource(tmp_path).discover()

    (tmp_path / "empty.pdf").unlink()
    (tmp_path / "bad.pdf").write_bytes(b"not a pdf")
    with pytest.raises(SourceValidationError):
        FilesystemDocumentSource(tmp_path).discover()

    (tmp_path / "bad.pdf").unlink()
    (tmp_path / "bad.docx").write_bytes(b"PK\x03\x04not-a-container")
    with pytest.raises(SourceValidationError):
        FilesystemDocumentSource(tmp_path).discover()

    (tmp_path / "bad.docx").unlink()
    with zipfile.ZipFile(tmp_path / "bomb.docx", "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr(
            "word/document.xml",
            '<!DOCTYPE x [<!ENTITY a "expanded">]><document>&a;</document>',
        )
    with pytest.raises(SourceValidationError):
        FilesystemDocumentSource(tmp_path).discover()


def test_text_pdf_preserves_page_boundaries(tmp_path: Path) -> None:
    raw = text_pdf("First page", "Second page")
    (tmp_path / "pages.pdf").write_bytes(raw)

    document = FilesystemDocumentSource(tmp_path).discover().documents[0]

    assert document.raw_bytes == raw
    assert document.media_type == "application/pdf"
    assert "First page" in document.content
    assert "Second page" in document.content
    assert "\f" in document.content


def test_structural_boundaries_token_caps_overlap_and_ids_are_deterministic() -> None:
    selected = parse_document_profiles(json.dumps(profiles()))
    profile = selected.profiles["generic"]
    paragraph = " ".join(f"token{i}" for i in range(400))
    text = (
        f"# First\n**ARTÍCULO 1.**\n{paragraph}\f# Second\n**ARTÍCULO 2.**\n{paragraph}"
    )
    first = build_chunks(text, document="law.md", checksum="a" * 64, profile=profile)
    second = build_chunks(text, document="law.md", checksum="a" * 64, profile=profile)
    encoding = tiktoken.get_encoding("o200k_base")

    assert first == second
    parents = [item for item in first if item.children_ids]
    leaves = [item for item in first if not item.children_ids]
    assert parents and leaves
    assert all(item.token_count <= 128 for item in parents)
    assert all(item.token_count <= 64 for item in leaves)
    assert {(item.page_number, item.article) for item in leaves} >= {
        (1, "1"),
        (2, "2"),
    }
    assert all(
        not ("ARTÍCULO 1" in item.text and "ARTÍCULO 2" in item.text) for item in leaves
    )
    by_parent: dict[str, list[object]] = {}
    for leaf in leaves:
        if leaf.parent_id:
            by_parent.setdefault(leaf.parent_id, []).append(leaf)
    for siblings in by_parent.values():
        ordered = sorted(siblings, key=lambda item: item.start_char)  # type: ignore[attr-defined]
        for left, right in zip(ordered, ordered[1:], strict=False):
            left_tokens = encoding.encode(left.text)  # type: ignore[attr-defined]
            right_tokens = encoding.encode(right.text)  # type: ignore[attr-defined]
            assert left_tokens[-8:] == right_tokens[:8]


def test_sqlite_migration_preserves_legacy_version_and_indexes_only_leaves(
    tmp_path: Path,
) -> None:
    path = tmp_path / "source.sqlite"
    connection = sqlite3.connect(path)
    connection.execute(
        """
        CREATE TABLE source_passages (
            id TEXT NOT NULL, document TEXT NOT NULL, article TEXT,
            paragraph TEXT, text TEXT NOT NULL, start_line INTEGER NOT NULL,
            end_line INTEGER NOT NULL, checksum TEXT NOT NULL,
            graph_version TEXT NOT NULL, PRIMARY KEY (id, graph_version)
        )
        """
    )
    connection.execute(
        """
        INSERT INTO source_passages VALUES
        ('legacy', 'old.md', NULL, NULL, 'old', 1, 1, ?, 'v0')
        """,
        ("a" * 64,),
    )
    connection.commit()
    connection.close()
    document = SimpleNamespace(
        relative_path="new.md",
        content="# Heading\n" + " ".join(f"value{i}" for i in range(300)),
        sha256="b" * 64,
        media_type="text/markdown",
        profile="generic",
    )
    configured = parse_document_profiles(json.dumps(profiles()))
    index = SourceIndex(path)
    index.rebuild_version("v1", [document], profiles=configured)

    assert index.has_version("v0")
    assert index.has_version("v1", configured.processing_fingerprint())
    connection = sqlite3.connect(path)
    parent_count = connection.execute(
        "SELECT count(*) FROM source_passages WHERE graph_version='v1' AND is_leaf=0"
    ).fetchone()[0]
    fts_count = connection.execute(
        "SELECT count(*) FROM source_passages_fts WHERE graph_version='v1'"
    ).fetchone()[0]
    leaf_count = connection.execute(
        "SELECT count(*) FROM source_passages WHERE graph_version='v1' AND is_leaf=1"
    ).fetchone()[0]
    connection.close()
    assert parent_count > 0
    assert fts_count == leaf_count
